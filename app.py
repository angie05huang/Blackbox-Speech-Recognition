"""
Main Flask application for BlackBox IT Application

Features:
- Transcribes uploaded audio files using WhisperX and outputs PDF/CSV/Excel
- Handles login and user management with project-based access control
- Supports project-specific file upload and retrieval using vectorstore-backed chatbot

Tech Stack:
- Flask, SQLAlchemy, WhisperX, HuggingFace, Chroma
"""
# Imports
from flask import Flask, request, send_file, jsonify, after_this_request, session
import os
import uuid
import io
import zipfile
import traceback
from docx import Document

import torch
import whisperx
import speech_recognition as sr
import openpyxl
import pandas as pd
from dotenv import load_dotenv
from pydub import AudioSegment
from flask_cors import CORS
from fpdf import FPDF
from openpyxl.styles import Alignment

from whisperx.diarize import DiarizationPipeline, assign_word_speakers
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash



from app_chroma_new import retrieve_bot_answer, delete_vector_store, get_all_files
from app_chroma_eat import chroma_ingest


# === Setup Flask and App ===
app = Flask(__name__)
CORS(app, supports_credentials=True, origins=["http://localhost:5173","http://localhost:5174", "http://172.65.32.103:3000"])
app.config['MAX_CONTENT_LENGTH'] = 2000 * 1024 * 1024  # 200 MB, for example
ALLOWED_EXTENSIONS = {'.wav', '.mp3', '.flac', '.aiff', '.aif', '.m4a', '.ogg', '.opus'}

# === Environment & Model Setup === 
load_dotenv()
token = os.getenv("HF_TOKEN") #Use Hugging Face Token in .env file
if not token:
    raise ValueError("Hugging Face token not found in environment.")

# === Whisper Model Setup === 
original_torch_load = torch.load
def trusted_torch_load(*args, **kwargs):
    kwargs['weights_only'] = False
    return original_torch_load(*args, **kwargs)
torch.load = trusted_torch_load
device = "cuda" if torch.cuda.is_available() else "cpu"
whisper_model = whisperx.load_model("base.en", device, compute_type="int8") #use base model

app.config['SQLALCHEMY_DATABASE_URI'] = "mysql+pymysql://root:blackbox123456@mysql:3306/blackbox_db"

#app.config['SQLALCHEMY_DATABASE_URI'] = "mysql+pymysql://root:blackbox123456@127.0.0.1:3306/blackbox_db"

app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Initialize database
db = SQLAlchemy(app) # SQL database

# === Database Models === 
class User(db.Model):
    """Database for user authentication and project linkage"""
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100))
    username = db.Column(db.String(100), unique=True, nullable=False)
    password_hash = db.Column(db.String(200), nullable=False)
    project_number = db.Column(db.String(100), nullable=False)
    Type = db.Column(db.String(50), nullable=False, default="Client") #default set to client

    def set_password(self, password):
        """Encrypts the password"""
        self.password_hash = generate_password_hash(password)
    
    def check_password(self, password):
        """Checks if the password is right"""
        return check_password_hash(self.password_hash, password)
    
    def get_project_number(self): 
        """Returns user project number"""
        return self.project_number

class ProjectFile(db.Model):
    """Tracks uploaded files linked to a specific project number"""
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(200), nullable=False)
    filepath = db.Column(db.String(300), nullable=False)  # Path to file saved
    project_number = db.Column(db.String(100), nullable=False)

    def __repr__(self):
        return f"<ProjectFile {self.filename} (Project {self.project_number})>"


# Create tables if not already created
with app.app_context():
    db.create_all()

# === Helper Functions === 
def get_extension(filename):
    """Returns filepath of specific file"""
    return os.path.splitext(filename)[1].lower()

def transcribe_whisperx(audio_path):
    """
    Transcribes audio using WhisperX and does speaker diarization (speaker labeling).

    Parameters:
        audio_path (str): Path to WAV audio
    Returns:
        List[Dict] with 'text' and 'speaker' fields, or error info
    """
    try:
        audio = whisperx.load_audio(audio_path)
        result = whisper_model.transcribe(audio)

        align_model, metadata = whisperx.load_align_model(language_code=result["language"], device=device)
        aligned_result = whisperx.align(result["segments"], align_model, metadata, audio, device)
        # Diarization for speaker labels
        try:
            diarize_pipeline = DiarizationPipeline(use_auth_token=token, device=device)
            diarize_segments = diarize_pipeline(audio)
            #result_with_speakers = whisperx.assign_word_speakers(diarize_segments, aligned_result["segments"])
            result_with_speakers = assign_word_speakers(diarize_segments, aligned_result)
            print("Diarization successful.")
        except Exception as diarize_err:
            print("Diarization failed:", diarize_err)
            for seg in aligned_result["segments"]:
                seg["speaker"] = "Speaker 1"
            result_with_speakers = aligned_result["segments"]
        print(aligned_result["segments"])
        return aligned_result["segments"]
    
    except Exception as e:
        tb = traceback.format_exc()
        print("WhisperX failed with error:", e)
        print(tb)
        # TEMPORARY DEBUG RETURN
        return {
            "error": f"WhisperX failed:\n{str(e)}",
            "traceback": tb
        }
    
def get_unique_user_names():
    """Returns list of all unique names in SQL database"""
    try:
        unique_names = (
            db.session.query(User.name)
            .filter(User.Type == "Client")
            .distinct()
            .order_by(User.name.asc())
            .all()
        )
        name_list = [name[0] for name in unique_names]
        return jsonify({"names": name_list}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    
# === Route to transcribe uploaded audio files ===
@app.route('/api/upload', methods=['POST'])
def upload_audio():
    """Transcribes audio file into text and returns text in pdf, csv, or excel file"""
    uploaded_files = request.files.getlist('files')
    print("Headers:", request.headers)
    print("Form keys:", request.form.keys())
    print("Files keys:", request.files.keys())
    uploaded_files = request.files.getlist('files')
    print(f"Number of files received: {len(uploaded_files)}")

    if not uploaded_files:
        return jsonify({'error': 'No files uploaded'}), 400

    aud_format = request.form.get('format', '').strip().lower()
    temp_paths = []
    output_files = []

    for audio_file in uploaded_files:
        
        ext = get_extension(audio_file.filename)
        if ext not in ALLOWED_EXTENSIONS:
            continue

        input_filename = f"{uuid.uuid4()}{ext}"
        audio_file.save(input_filename)
        base_name = os.path.splitext(audio_file.filename)[0]
        temp_paths.append(input_filename)

        # Converts the audio file to wav if not wav
        
        try:
            audio = AudioSegment.from_file(input_filename)
            audio = audio.set_channels(1).set_frame_rate(16000)
            processed_filename = f"{uuid.uuid4()}.wav"
            audio.export(processed_filename, format="wav")
            temp_paths.append(processed_filename)
        except Exception as e:
            return jsonify({"error": f"Could not process {audio_file.filename}: {str(e)}"}), 500

        # Segments the audio file into chunks for whisperx to transcribe
        
        segments = transcribe_whisperx(processed_filename)
        if not segments or not isinstance(segments, list):
            if isinstance(segments, dict) and "error" in segments:
                return jsonify(segments), 500  # show actual error
            return jsonify({"error": "WhisperX transcription failed."}), 500
        # Transcribe to csv
        csv_data = [[seg.get("speaker", "Unknown"), seg["text"].strip()] for seg in segments if isinstance(seg, dict) and seg.get("text")]
        if aud_format == 'csv':
            csv_filename = f"{uuid.uuid4()}.csv"
            
            # Format speaker + cleaned transcript
            csv_data = [
                [seg.get("speaker", f"Speaker {i+1}"), seg["text"].replace('. ', '.\n').strip()]
                for i, seg in enumerate(segments) if seg.get("text")
            ]


            df = pd.DataFrame(csv_data, columns=["Speaker", "Transcript"])
            df.to_csv(csv_filename, index=False)

            output_files.append((csv_filename, f"{base_name}.csv"))
            temp_paths.append(csv_filename)

        
        # Transcribe to excel
        elif aud_format == 'xlsx':
            excel_filename = f"{uuid.uuid4()}.xlsx"
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.append(["Speaker", "Transcript"])

            for seg in segments:
                if seg.get("text"):
                    speaker = seg.get("speaker", "Unknown")
                    line = seg["text"].strip()
                    ws.append([speaker, line])

            for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=2, max_col=2):
                for cell in row:
                    cell.alignment = Alignment(wrap_text=True)

            ws.column_dimensions['A'].width = 20
            ws.column_dimensions['B'].width = 80

            wb.save(excel_filename)
            output_files.append((excel_filename, f"{base_name}.xlsx"))
            temp_paths.append(excel_filename)

        # Transcribe to pdf
        elif aud_format == 'pdf':
            pdf_filename = f"{uuid.uuid4()}.pdf"
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.set_font("Arial", size=12)

            for seg in segments:
                if seg.get("text"):
                    speaker = seg.get("speaker", "Unknown")
                    line = seg["text"].strip()
                    pdf.multi_cell(0, 10, txt=f"{speaker}: {line}", align='L')
                    pdf.ln(1)

            pdf.output(pdf_filename)
            output_files.append((pdf_filename, f"{base_name}.pdf"))
            temp_paths.append(pdf_filename)
        
        elif aud_format == 'docx':
            docx_filename = f"{uuid.uuid4()}.docx"
            document = Document()
            document.add_heading('Transcription Output', level=1)

            for seg in segments:
                if seg.get("text"):
                    speaker = seg.get("speaker", "Unknown")
                    line = seg["text"].strip()
                    document.add_paragraph(f"{speaker}: {line}")

            document.save(docx_filename)
            output_files.append((docx_filename, f"{base_name}.docx"))
            temp_paths.append(docx_filename)


        else:
            return jsonify({"error": f"Unsupported format: {aud_format}"}), 400

        # Remove original uploaded file
        if ext != '.wav':
            os.remove(input_filename)
    
    # Returns a zip for users to click into
    zip_io = io.BytesIO()
    with zipfile.ZipFile(zip_io, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for path, arcname in output_files:
            zipf.write(path, arcname=arcname)

    # Cleans up temporary files
    @after_this_request
    def cleanup(response):
        for path in temp_paths:
            try:
                os.remove(path)
            except FileNotFoundError:
                pass
        return response

    zip_io.seek(0)
    return send_file(
        zip_io,
        as_attachment=True,
        download_name="translation_bundle.zip",
        mimetype='application/zip'
    )

# === Route to login authentication ===
@app.route('/api/login', methods=['POST'])
def login():
    """Retrieves username and password from frontend, checks if they match"""
    data = request.form
    username = data.get('username')
    password = data.get('password')

    print(f"Login attempt: username={username}, password={password}")
    
    user = User.query.filter_by(username=username).first()
    if user and user.check_password(password):
        #session['user'] = username
        return jsonify({"Loggedin": True, "project_number": user.project_number, "Type": user.Type, "Name":user.name}), 200

    # Returns 401 if incorrect
    return jsonify({"success": False, "message": "Invalid credentials"}), 401

# === Route to add_user ===
@app.route('/api/add_user', methods=['POST'])
def add_user():
    """Creates new user with their login and project information"""
    data = request.form
    name = data.get('name')
    username = data.get('username')
    password = data.get('password')

    # Checks if user exists already
    if User.query.filter_by(username=username).first():
        return jsonify({"success": False, "message": "User already exists"}), 400

    project_number = data.get('project_number')
    user_type = data.get('type', 'Client')  # Default: 'Client'

    new_user = User(name=name, username=username, project_number=project_number, Type=user_type)
    new_user.set_password(password)
    db.session.add(new_user)
    db.session.commit()
    return jsonify({"success": True, "message": "User added successfully"}), 201

@app.route('/api/modify_user', methods=['POST'])
def ModifyUser():
    """Updates user fields: name, username, password, type, project_number"""
    data = request.form

    # Admin check
    user_type = data.get('Type')
    if user_type != 'Admin':
        return jsonify({"error": "Unauthorized"}), 403

    user_id = data.get('id')
    if not user_id:
        return jsonify({"error": "Missing user ID"}), 400

    user = User.query.get(user_id)
    if not user:
        return jsonify({"error": "User not found"}), 404

    # Update only if provided
    if 'name' in data:
        user.name = data.get('name')
    if 'username' in data:
        user.username = data.get('username')
    if 'project_number' in data:
        user.project_number = data.get('project_number')
    if 'user_type' in data:
        user.Type = data.get('user_type')  # field is 'Type' in DB, 'user_type' in frontend
    if 'password' in data and data.get('password').strip() != '':
        user.set_password(data.get('password'))
    
    if data.get('delete_user', '').lower() == 'true':
        try:
            db.session.delete(user)
            db.session.commit()
            return jsonify({"message": f"User {user.name} deleted successfully"}), 200
        except Exception as e:
            db.session.rollback()
            return jsonify({"error": f"Failed to delete user: {str(e)}"}), 500


    db.session.commit()
    return jsonify({"success": True, "message": "User updated successfully"}), 200

# === Route to check all users ===
@app.route('/api/allusers', methods=['GET'])
def get_allusers():
    """Returns information of users in SQL database"""
    allusers = User.query.all()
    return jsonify([
        {"id": user.id, "name": user.name, "username": user.username, "project_number": user.project_number, "user_type": user.Type}
        for user in allusers
    ])

# === Route to check users ===
@app.route('/api/users', methods=['GET'])
def get_users():
    """Returns information of users in SQL database"""
    users = User.query.filter_by(Type="Client").all()
    return jsonify([
        {"id": user.id, "name": user.name, "username": user.username, "project_number": user.project_number, "user_type": user.Type}
        for user in users
    ])
    #return {"users":[{"id": 0, "name": "Alexa", "username": "Ale", "project_number": 0}, {"id": 1, "name": "Alex", "username": "Al", "project_number": 1},{"id": 2, "name": "Angie", "username": "Ange", "project_number": 1}]},200
# === Route to get all docs ==
@app.route('/api/doclist', methods = ['POST'])
def docs():
    project_number = request.json.get("project_number")
    return {"documents":get_all_files(project_number) },200



# === Route to chatbox ===
chat_history = []
@app.route('/api/llama2', methods = ['POST'])
def talktalk():
    """Route for chat messages to project-specific vectorstore chatbot"""
    print("Incoming JSON:", request.json)
    message = request.json["message"] # Retrieves message from frontend
    project_number = request.json.get("project_number") # Retrieves project number from frontend

    if not project_number:
        print("project number not here")
        return jsonify({"error": "Missing project number"}), 400
    
    print("Calling retrieve_bot_answer with:", message, project_number)
    
    # Calls method sendMessage which takes in message and project_number as parameters 
    answer= sendMessage(message, project_number)
    return jsonify({"ChatBot": answer}), 200

def sendMessage(message, project_number):
    """Calls helper function retrieve_bot_answer from app_chroma_new file"""
    answer = retrieve_bot_answer(message, project_number)
    print(answer)
    return answer

# === Route to upload files into chat ===
@app.route('/api/uploadllama2', methods=["POST"])
def uploadLlamaChroma():
    """Handles file upload for ingestion into project-specific vectorstore"""
    print("Received request:")
    print("Files:", request.files)
    print("Form:", request.form)
    project_number = request.form.get("project_number")
    print("Project number:", project_number)
    uploaded_files = request.files.getlist('files')

    if not project_number:
        return jsonify({'error': 'Missing project number'}), 400
    delete_vector_store(project_number) #deletes all old files in vectorstore
    for file in uploaded_files:
        save_and_ingest_file(file, project_number)

    return {}, 200

def save_and_ingest_file(file, project_number):
    """Save uploaded file and ingest into corresponding vectorstore"""
    # Save file locally
    

    filename = f"{file.filename}"
    save_path = os.path.join("project_files", filename)
    os.makedirs("project_files", exist_ok=True)
    file.save(save_path)

    # Add file to database
    new_entry = ProjectFile(filename=file.filename, filepath=save_path, project_number=project_number)
    db.session.add(new_entry)
    db.session.commit()

    # Ingest into vectorstore with directory per project using function from app_chroma_eat file
    chroma_ingest(save_path, vector_dir=f"./vector_store/vectorstore_{project_number}")

if __name__ == "__main__":
    app.run(host="0.0.0.0", debug = True)
